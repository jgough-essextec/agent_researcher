from docx import Document
from docx.shared import Inches
from .base import BaseReportSection


class InfrastructureMapsSection(BaseReportSection):
    def render(self, document: Document) -> None:
        document.add_page_break()
        document.add_heading('Appendix C: Infrastructure Maps', level=1)

        screenshots = self.job.screenshots.all()
        if not screenshots.exists():
            document.add_paragraph("No infrastructure screenshots were captured during this assessment.")
            return

        for screenshot in screenshots:
            try:
                document.add_picture(screenshot.image.path, width=Inches(6))
                document.add_paragraph(
                    f"{screenshot.source.title()} — {screenshot.caption or 'No caption'}"
                )
            except Exception:
                document.add_paragraph(f"[Screenshot not available: {screenshot.source}]")
