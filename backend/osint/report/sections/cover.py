from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import BaseReportSection


class CoverSection(BaseReportSection):
    def render(self, document: Document) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.bold = True

        document.add_paragraph()

        title = document.add_heading('External Attack Surface &', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title2 = document.add_heading('Threat Intelligence Assessment', level=0)
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()

        prospect_p = document.add_paragraph()
        prospect_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = prospect_p.add_run(self.job.organization_name)
        run.font.size = Pt(20)
        run.bold = True

        document.add_paragraph()

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Prepared by: Your Security Practice\n")
        meta.add_run(f"Primary Domain: {self.job.primary_domain}\n")
        meta.add_run(f"Date: {date.today().strftime('%B %d, %Y')}\n")
        meta.add_run(
            "\nAll intelligence gathered using passive, legal OSINT only.\n"
            "No active scanning, no exploitation, no authenticated access."
        )

        document.add_page_break()
